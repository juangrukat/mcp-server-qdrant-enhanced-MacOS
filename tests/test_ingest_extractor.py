import json

from mcp_server_qdrant.ingest import extractor
from mcp_server_qdrant.ingest.extractor import (
    SUPPORTED_EXTENSIONS,
    ExtractedDocument,
    PdfProfile,
    build_chunks,
    extract_text,
)


def test_json_is_rendered_as_searchable_path_value_text(tmp_path):
    path = tmp_path / "sample.json"
    path.write_text(
        json.dumps({"name": "Ada", "skills": ["math", "systems"], "active": True}),
        encoding="utf-8",
    )

    doc = extract_text(str(path))

    assert doc.error is None
    assert doc.extractor_used == "json"
    assert "name: Ada" in doc.text
    assert "skills[0]: math" in doc.text
    assert "active: True" in doc.text


def test_tsv_is_rendered_as_row_column_text(tmp_path):
    path = tmp_path / "sample.tsv"
    path.write_text("name\trole\nGrace\tcompiler pioneer\n", encoding="utf-8")

    doc = extract_text(str(path))

    assert doc.error is None
    assert doc.extractor_used == "tsv"
    assert "Row 1" in doc.text
    assert "name: Grace" in doc.text
    assert "role: compiler pioneer" in doc.text


def test_plain_textish_extensions_are_supported(tmp_path):
    path = tmp_path / "settings.yaml"
    path.write_text("service: qdrant\nmode: local\n", encoding="utf-8")

    doc = extract_text(str(path))

    assert ".yaml" in SUPPORTED_EXTENSIONS
    assert doc.error is None
    assert doc.extractor_used == "plain-utf-8"
    assert "service: qdrant" in doc.text


def test_build_chunks_uses_conservative_env_chunk_size(monkeypatch):
    monkeypatch.setenv("QDRANT_INGEST_CHUNK_SIZE", "20")
    monkeypatch.setenv("QDRANT_INGEST_CHUNK_OVERLAP", "5")
    doc = ExtractedDocument(
        text="abcdefghijklmnopqrstuvwxyz",
        extractor_used="test",
        char_count=26,
    )

    chunks = build_chunks(doc, {"source": "unit"})

    assert [chunk.text for chunk in chunks] == ["abcdefghijklmnopqrst", "pqrstuvwxyz"]
    assert chunks[0].metadata["total_chunks"] == 2


def test_build_chunks_normalizes_internal_whitespace():
    """PDF extractors produce double-spaced text from column layouts.
    build_chunks should collapse multiple spaces to one so phrase matching works."""
    doc = ExtractedDocument(
        text="Critical  reading,  critical  thinking,  discussion  skills.",
        extractor_used="test",
        char_count=60,
    )

    chunks = build_chunks(doc, {"source": "unit"})

    assert len(chunks) == 1
    assert "Critical reading, critical thinking, discussion skills." in chunks[0].text


def test_build_chunks_normalizes_tabs_but_preserves_newlines():
    """Tabs are collapsed to spaces; paragraph-separating newlines are kept."""
    doc = ExtractedDocument(
        text="First\tparagraph.\n\nSecond\t\tparagraph.",
        extractor_used="test",
        char_count=38,
    )

    chunks = build_chunks(doc, {"source": "unit"})

    # Both paragraphs should be present and tabs collapsed
    combined = " ".join(c.text for c in chunks)
    assert "First paragraph." in combined
    assert "Second paragraph." in combined


def test_docx_tables_are_extracted_in_document_order(tmp_path):
    from docx import Document

    path = tmp_path / "table.docx"
    source = Document()
    source.add_paragraph("Before table")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Ada"
    table.cell(1, 1).text = "Mathematician"
    source.add_paragraph("After table")
    source.save(path)

    doc = extract_text(str(path))

    assert doc.error is None
    assert doc.extractor_used == "python-docx"
    assert doc.text.index("Before table") < doc.text.index("[TABLE]")
    assert "Row 1 | column_1: Name | column_2: Role" in doc.text
    assert "Row 2 | column_1: Ada | column_2: Mathematician" in doc.text
    assert doc.text.index("[/TABLE]") < doc.text.index("After table")


def test_scanned_pdf_preflight_skips_full_text_extractors(monkeypatch, tmp_path):
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-1.4\n")

    monkeypatch.setattr(
        extractor,
        "profile_pdf",
        lambda _: PdfProfile(
            has_text=False,
            has_images=True,
            is_probably_scanned=True,
            is_encrypted=False,
            requires_password=False,
            pages_with_text=[],
            pages_sampled=1,
            page_count=1,
        ),
    )

    def fail_if_called(_):
        raise AssertionError("full PDF extraction should not run for scanned preflight")

    monkeypatch.setattr(extractor, "_pdf_pdfminer", fail_if_called)
    monkeypatch.setattr(extractor, "_pdf_pypdf", fail_if_called)

    doc = extract_text(str(path))

    assert doc.extractor_used == "pdf-preflight"
    assert doc.page_count == 1
    assert doc.char_count == 0
    assert "OCR is not implemented" in doc.error


def test_password_protected_pdf_fails_before_extraction(monkeypatch):
    monkeypatch.setattr(
        extractor,
        "profile_pdf",
        lambda _: PdfProfile(
            has_text=False,
            has_images=False,
            is_probably_scanned=False,
            is_encrypted=True,
            requires_password=True,
            pages_with_text=[],
            pages_sampled=0,
            page_count=0,
        ),
    )

    def fail_if_called(_):
        raise AssertionError("encrypted PDF should not be extracted")

    monkeypatch.setattr(extractor, "_pdf_pdfminer", fail_if_called)
    monkeypatch.setattr(extractor, "_pdf_pypdf", fail_if_called)

    doc = extractor._extract_pdf("/tmp/encrypted.pdf")

    assert doc.extractor_used == "pdf-preflight"
    assert doc.char_count == 0
    assert "password-protected" in doc.error
