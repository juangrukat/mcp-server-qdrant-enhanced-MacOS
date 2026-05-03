"""
Text extraction for ingest pipeline.
Priority per format:
  .txt / .md  → direct UTF-8 read (charset fallback)
  structured  → JSON/JSONL/CSV/TSV rendered into searchable text
  .pdf        → pdfminer.six primary, pypdf fallback
  .docx       → python-docx
"""

import csv
import json
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

PLAIN_TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".text",
    ".conf", ".cfg", ".ini", ".env",
    ".yaml", ".yml", ".toml", ".xml", ".html", ".htm",
    ".css", ".scss", ".less",
    ".js", ".jsx", ".ts", ".tsx", ".py", ".java", ".c", ".cc", ".cpp",
    ".h", ".hpp", ".cs", ".go", ".rs", ".rb", ".php", ".swift",
    ".kt", ".kts", ".sh", ".bash", ".zsh", ".fish", ".sql",
    ".graphql", ".gql",
}
STRUCTURED_TEXT_EXTENSIONS = {".json", ".jsonl", ".csv", ".tsv"}
SUPPORTED_EXTENSIONS = PLAIN_TEXT_EXTENSIONS | STRUCTURED_TEXT_EXTENSIONS | {".pdf", ".docx"}

# Max chars per chunk; overlap in chars. Qwen3/Candle is sensitive to inputs
# that exceed its tokenizer split limits, so keep local chunks conservative.
DEFAULT_CHUNK_SIZE = 700
DEFAULT_CHUNK_OVERLAP = 70


@dataclass
class ExtractedDocument:
    text: str
    extractor_used: str
    char_count: int
    page_count: int | None = None
    error: str | None = None


@dataclass
class PdfProfile:
    has_text: bool
    has_images: bool
    is_probably_scanned: bool
    is_encrypted: bool
    requires_password: bool
    pages_with_text: list[int]
    pages_sampled: int
    page_count: int
    error: str | None = None


@dataclass
class Chunk:
    text: str
    chunk_index: int
    total_chunks: int
    metadata: dict = field(default_factory=dict)


def extract_text(path: str) -> ExtractedDocument:
    """Dispatch to the right extractor based on file extension."""
    ext = Path(path).suffix.lower()
    if ext in PLAIN_TEXT_EXTENSIONS:
        return _extract_plain(path)
    if ext == ".json":
        return _extract_json(path)
    if ext == ".jsonl":
        return _extract_jsonl(path)
    if ext in (".csv", ".tsv"):
        return _extract_delimited(path, delimiter="\t" if ext == ".tsv" else ",")
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    return ExtractedDocument(
        text="",
        extractor_used="none",
        char_count=0,
        error=f"Unsupported file type: {ext}",
    )


def build_chunks(doc: ExtractedDocument, file_metadata: dict) -> list[Chunk]:
    """
    Split extracted text into overlapping chunks.
    Splits on paragraph boundaries first, then hard-cuts long paragraphs.
    """
    chunk_size = _env_int("QDRANT_INGEST_CHUNK_SIZE", DEFAULT_CHUNK_SIZE)
    chunk_overlap = min(
        _env_int("QDRANT_INGEST_CHUNK_OVERLAP", DEFAULT_CHUNK_OVERLAP),
        max(0, chunk_size - 1),
    )
    stride = max(1, chunk_size - chunk_overlap)
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", doc.text) if p.strip()]
    raw_chunks: list[str] = []
    current = ""

    for para in paragraphs:
        if len(para) > chunk_size:
            # Hard split the oversized paragraph
            if current:
                raw_chunks.append(current)
                current = ""
            for i in range(0, len(para), stride):
                raw_chunks.append(para[i : i + chunk_size])
        elif len(current) + len(para) + 2 > chunk_size:
            raw_chunks.append(current)
            current = para
        else:
            current = (current + "\n\n" + para).strip() if current else para

    if current:
        raw_chunks.append(current)

    total = len(raw_chunks)
    chunks = []
    for i, text in enumerate(raw_chunks):
        chunk_meta = {**file_metadata, "chunk_index": i, "total_chunks": total}
        chunks.append(Chunk(text=text, chunk_index=i, total_chunks=total, metadata=chunk_meta))
    return chunks


def _env_int(name: str, default: int) -> int:
    raw = os.getenv(name)
    if not raw:
        return default
    try:
        return max(1, int(raw))
    except ValueError:
        logger.warning("Invalid %s=%r; using %s", name, raw, default)
        return default


# --- Plain text ---

def _extract_plain(path: str) -> ExtractedDocument:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = Path(path).read_text(encoding=enc)
            # Strip YAML/TOML frontmatter from markdown
            if Path(path).suffix.lower() in (".md", ".markdown"):
                text = _strip_frontmatter(text)
            return ExtractedDocument(
                text=text,
                extractor_used=f"plain-{enc}",
                char_count=len(text),
            )
        except UnicodeDecodeError:
            continue
        except Exception as e:
            return ExtractedDocument(text="", extractor_used="plain", char_count=0, error=str(e))
    return ExtractedDocument(text="", extractor_used="plain", char_count=0, error="Could not decode file")


def _strip_frontmatter(text: str) -> str:
    """Remove YAML (---) or TOML (+++) frontmatter from markdown."""
    for fence in ("---", "+++"):
        if text.startswith(fence):
            end = text.find(fence, len(fence))
            if end != -1:
                return text[end + len(fence):].lstrip("\n")
    return text


# --- Structured text ---

def _extract_json(path: str) -> ExtractedDocument:
    try:
        raw = Path(path).read_text(encoding="utf-8")
        data = json.loads(raw)
        text = _json_to_searchable_text(data)
        return ExtractedDocument(
            text=text,
            extractor_used="json",
            char_count=len(text),
        )
    except UnicodeDecodeError:
        return _extract_plain(path)
    except Exception as e:
        return ExtractedDocument(text="", extractor_used="json", char_count=0, error=str(e))


def _extract_jsonl(path: str) -> ExtractedDocument:
    try:
        rendered = []
        for line_no, line in enumerate(Path(path).read_text(encoding="utf-8").splitlines(), start=1):
            if not line.strip():
                continue
            data = json.loads(line)
            rendered.append(f"Record {line_no}\n{_json_to_searchable_text(data)}")
        text = "\n\n".join(rendered)
        return ExtractedDocument(
            text=text,
            extractor_used="jsonl",
            char_count=len(text),
        )
    except UnicodeDecodeError:
        return _extract_plain(path)
    except Exception as e:
        return ExtractedDocument(text="", extractor_used="jsonl", char_count=0, error=str(e))


def _extract_delimited(path: str, *, delimiter: str) -> ExtractedDocument:
    extractor = "tsv" if delimiter == "\t" else "csv"
    try:
        text = Path(path).read_text(encoding="utf-8-sig")
        reader = csv.reader(text.splitlines(), delimiter=delimiter)
        rows = list(reader)
        if not rows:
            return ExtractedDocument(text="", extractor_used=extractor, char_count=0)

        header = [cell.strip() for cell in rows[0]]
        has_header = any(header)
        rendered_rows = []
        for row_index, row in enumerate(rows[1:] if has_header else rows, start=1):
            cells = []
            for col_index, value in enumerate(row):
                label = header[col_index] if has_header and col_index < len(header) and header[col_index] else f"column_{col_index + 1}"
                if value.strip():
                    cells.append(f"{label}: {value.strip()}")
            if cells:
                rendered_rows.append(f"Row {row_index}\n" + "\n".join(cells))

        text_out = "\n\n".join(rendered_rows)
        return ExtractedDocument(
            text=text_out,
            extractor_used=extractor,
            char_count=len(text_out),
        )
    except UnicodeDecodeError:
        return _extract_plain(path)
    except Exception as e:
        return ExtractedDocument(text="", extractor_used=extractor, char_count=0, error=str(e))


def _json_to_searchable_text(value: Any, *, prefix: str = "") -> str:
    lines = []
    if isinstance(value, dict):
        for key, child in value.items():
            child_prefix = f"{prefix}.{key}" if prefix else str(key)
            lines.append(_json_to_searchable_text(child, prefix=child_prefix))
    elif isinstance(value, list):
        for index, child in enumerate(value):
            child_prefix = f"{prefix}[{index}]" if prefix else f"[{index}]"
            lines.append(_json_to_searchable_text(child, prefix=child_prefix))
    elif value is not None:
        label = prefix or "value"
        lines.append(f"{label}: {value}")
    return "\n".join(line for line in lines if line)


# --- PDF ---

def _extract_pdf(path: str) -> ExtractedDocument:
    profile = profile_pdf(path)
    if profile.requires_password:
        return ExtractedDocument(
            text="",
            extractor_used="pdf-preflight",
            char_count=0,
            page_count=profile.page_count,
            error="PDF is encrypted/password-protected and cannot be ingested without a password.",
        )
    if profile.is_probably_scanned:
        return ExtractedDocument(
            text="",
            extractor_used="pdf-preflight",
            char_count=0,
            page_count=profile.page_count,
            error=(
                "PDF appears to be image-only/scanned with no embedded text layer "
                f"in the first {profile.pages_sampled} sampled page(s). OCR is not implemented."
            ),
        )

    text, pages, extractor = _pdf_pdfminer(path)
    if not text:
        text, pages, extractor = _pdf_pypdf(path)
    if not text:
        return ExtractedDocument(text="", extractor_used="pdf-failed", char_count=0, error="All PDF extractors returned empty text")
    return ExtractedDocument(text=text, extractor_used=extractor, char_count=len(text), page_count=pages)


def profile_pdf(path: str, *, sample_pages: int = 5, text_threshold: int = 16) -> PdfProfile:
    """
    Cheaply classify whether a PDF likely has a text layer before full extraction.
    This is not OCR; it only samples existing PDF structure and extractable text.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        if reader.is_encrypted:
            try:
                decrypt_result = reader.decrypt("")
            except Exception:
                decrypt_result = 0
            if not decrypt_result:
                return PdfProfile(
                    has_text=False,
                    has_images=False,
                    is_probably_scanned=False,
                    is_encrypted=True,
                    requires_password=True,
                    pages_with_text=[],
                    pages_sampled=0,
                    page_count=0,
                    error="PDF is encrypted/password-protected.",
                )

        page_count = len(reader.pages)
        pages_to_sample = min(sample_pages, page_count)
        has_images = False
        pages_with_text: list[int] = []

        for index in range(pages_to_sample):
            page = reader.pages[index]
            fonts, images = _pdf_page_resource_flags(page)
            has_images = has_images or images

            text = (page.extract_text() or "").strip()
            if len(text) >= text_threshold:
                pages_with_text.append(index)
                break

            if fonts:
                continue

            miner_text = _pdfminer_page_text(path, index)
            if len(miner_text.strip()) >= text_threshold:
                pages_with_text.append(index)
                break

        has_text = bool(pages_with_text)
        return PdfProfile(
            has_text=has_text,
            has_images=has_images,
            is_probably_scanned=has_images and not has_text and pages_to_sample > 0,
            is_encrypted=reader.is_encrypted,
            requires_password=False,
            pages_with_text=pages_with_text,
            pages_sampled=pages_to_sample,
            page_count=page_count,
        )
    except Exception as e:
        logger.debug(f"PDF preflight failed for {path}: {e}")
        return PdfProfile(
            has_text=False,
            has_images=False,
            is_probably_scanned=False,
            is_encrypted=False,
            requires_password=False,
            pages_with_text=[],
            pages_sampled=0,
            page_count=0,
            error=str(e),
        )


def _pdf_page_resource_flags(page: Any) -> tuple[bool, bool]:
    resources = _pdf_resolve_object(page.get("/Resources", {})) or {}
    fonts = bool(resources.get("/Font"))
    xobjects = _pdf_resolve_object(resources.get("/XObject", {})) or {}
    has_images = False
    for obj in xobjects.values():
        resolved = _pdf_resolve_object(obj)
        if resolved and resolved.get("/Subtype") == "/Image":
            has_images = True
            break
    return fonts, has_images


def _pdf_resolve_object(obj: Any) -> Any:
    if hasattr(obj, "get_object"):
        return obj.get_object()
    return obj


def _pdfminer_page_text(path: str, page_number: int) -> str:
    try:
        from pdfminer.high_level import extract_text as pm_extract

        return pm_extract(path, page_numbers=[page_number]) or ""
    except Exception as e:
        logger.debug(f"pdfminer page preflight failed for {path} page {page_number}: {e}")
        return ""


def _pdf_pdfminer(path: str) -> tuple[str, int | None, str]:
    try:
        from pdfminer.high_level import extract_text as pm_extract
        from pdfminer.high_level import extract_pages
        text = pm_extract(path)
        pages = sum(1 for _ in extract_pages(path))
        return (text or "").strip(), pages, "pdfminer"
    except Exception as e:
        logger.debug(f"pdfminer failed for {path}: {e}")
        return "", None, "pdfminer"


def _pdf_pypdf(path: str) -> tuple[str, int | None, str]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = len(reader.pages)
        text = "\n\n".join(
            page.extract_text() or "" for page in reader.pages
        ).strip()
        return text, pages, "pypdf"
    except Exception as e:
        logger.debug(f"pypdf failed for {path}: {e}")
        return "", None, "pypdf"


# --- DOCX ---

def _extract_docx(path: str) -> ExtractedDocument:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(path)
        body_parts = _render_docx_container(doc, Paragraph, Table)
        header_footer_parts = _render_docx_headers_footers(doc, Paragraph, Table)
        text = "\n\n".join(part for part in [body_parts, header_footer_parts] if part)
        return ExtractedDocument(
            text=text,
            extractor_used="python-docx",
            char_count=len(text),
        )
    except Exception as e:
        return ExtractedDocument(text="", extractor_used="python-docx", char_count=0, error=str(e))


def _render_docx_container(container: Any, paragraph_type: type, table_type: type) -> str:
    parts = []
    for block in container.iter_inner_content():
        if isinstance(block, paragraph_type):
            text = block.text.strip()
            if text:
                parts.append(text)
        elif isinstance(block, table_type):
            table_text = _render_docx_table(block, paragraph_type, table_type)
            if table_text:
                parts.append("[TABLE]\n" + table_text + "\n[/TABLE]")
    return "\n\n".join(parts)


def _render_docx_table(table: Any, paragraph_type: type, table_type: type) -> str:
    lines = []
    for row_index, row in enumerate(table.rows, start=1):
        cells = []
        for col_index, cell in enumerate(row.cells, start=1):
            cell_text = _render_docx_container(cell, paragraph_type, table_type).replace("\n", " ").strip()
            if cell_text:
                cells.append(f"column_{col_index}: {cell_text}")
        if cells:
            lines.append(f"Row {row_index} | " + " | ".join(cells))
    return "\n".join(lines)


def _render_docx_headers_footers(doc: Any, paragraph_type: type, table_type: type) -> str:
    parts = []
    seen: set[int] = set()
    for section_index, section in enumerate(doc.sections, start=1):
        for label, container in (
            ("header", section.header),
            ("footer", section.footer),
            ("first_header", section.first_page_header),
            ("first_footer", section.first_page_footer),
            ("even_header", section.even_page_header),
            ("even_footer", section.even_page_footer),
        ):
            container_id = id(container._element)
            if container_id in seen:
                continue
            seen.add(container_id)
            text = _render_docx_container(container, paragraph_type, table_type)
            if text:
                parts.append(f"[{label.upper()} {section_index}]\n{text}")
    return "\n\n".join(parts)
